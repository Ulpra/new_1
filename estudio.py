from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches

# Definir la fecha de inicio
start_date = datetime(2025, 6, 30)

# Crear el documento
doc = Document()
doc.add_heading("Calendario de Estudio - Oposición", level=1)
doc.add_paragraph("Inicio: Lunes 30 de Junio de 2025\nDuración: 12 semanas\nObjetivo: Estudiar 100 temas (aprox. 8-9 por semana)")

# Definir los bloques diarios
daily_schedule = [
    ("16:30 – 18:30", "Estudio"),
    ("18:30 – 20:00", "Esquema/Resumen/Test")
]

# Generar las 12 semanas
topics_per_week = [8] * 12
extra_topics = 100 - sum(topics_per_week)
for i in range(extra_topics):
    topics_per_week[i] += 1  # Distribuir los temas restantes

topic_counter = 1
for week in range(12):
    week_start = start_date + timedelta(weeks=week)
    week_end = week_start + timedelta(days=6)
    doc.add_heading(f"Semana {week + 1}: {week_start.strftime('%d %b')} – {week_end.strftime('%d %b')}", level=2)
    doc.add_paragraph(f"Temas: {topic_counter} – {topic_counter + topics_per_week[week] - 1}")
    
    for day in range(7):
        current_day = week_start + timedelta(days=day)
        doc.add_heading(current_day.strftime('%A %d %B'), level=3)
        # Plan del día: alternar estudio y esquema/test para cada tema
        topics_today = 2 if topic_counter + 1 <= 100 else 1
        if day == 5:  # Sábado: repaso + test general
            doc.add_paragraph("09:00 – 11:00: Repaso errores y test semanales")
            doc.add_paragraph("11:00 – 13:00: Test general (20 preguntas)")
        elif day == 6:  # Domingo: repaso + revisión esquemas
            doc.add_paragraph("09:00 – 11:00: Revisión test + esquemas")
            doc.add_paragraph("11:00 – 13:00: Repaso de esquemas de la semana")
        else:
            for block in daily_schedule:
                if topic_counter <= 100:
                    doc.add_paragraph(f"{block[0]}: {block[1]} Tema {topic_counter}")
                    topic_counter += 1
                else:
                    doc.add_paragraph(f"{block[0]}: Repaso general")
                    
doc_path = "/mnt/data/Calendario_Estudio_Oposicion_12_Semanas.docx"
doc.save(doc_path)
doc_path
